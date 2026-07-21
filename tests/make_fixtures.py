#!/usr/bin/env python3
"""Generate the Featherpress docx edge-case fixtures.

This is the shipped, version-controlled source of truth for the test
fixtures. The .docx binaries it emits are generated on demand into
fixtures/ (git-ignored), so reviewers read this generator, not opaque
zip bytes.

edges.docx exercises all three docx edge classes ISS-002 covers, in one
file: a table, an inline image, and a footnote reference (the footnote
Part is injected directly because python-docx 1.2.0 has no footnote
author API).

All content is ASCII, no em dashes, deterministic strings so the probes
can assert exact substrings. cp1252 default locale on Windows: every
file IO passes encoding explicitly.
"""

import io
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.opc.constants import RELATIONSHIP_TYPE, CONTENT_TYPE

SCRIPT_DIR = Path(__file__).resolve().parent
FIXTURES_DIR = SCRIPT_DIR / "fixtures"

def _png_1x1() -> bytes:
    """Build a valid, deterministic 1x1 RGB PNG with correct chunk CRCs.

    Hand-authored byte literals proved fragile (python-docx's PNG parser
    walks chunk offsets and needs correct lengths/CRCs), so construct the
    chunks programmatically. zlib output is deterministic, so the bytes are
    stable across runs.
    """
    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return struct.pack(">I", len(data)) + body + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1, 8-bit, RGB
    raw = b"\x00\x00\x00\x00"                             # filter 0 + one RGB pixel
    idat = zlib.compress(raw, 9)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


PNG_1X1 = _png_1x1()

# The footnote body deliberately contains an ampersand (&amp; in the XML,
# reading back as &) to pin the escaping fix in parse_docx: an unescaped &
# would corrupt HTML/EPUB and crash reportlab's Paragraph parser in the PDF
# build. The separator / continuationSeparator entries mirror real Word and
# exercise the @w:type skip.
FOOTNOTES_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:footnote w:type="separator" w:id="-1"><w:p/></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p/></w:footnote>'
    '<w:footnote w:id="2"><w:p><w:r><w:t>Cited source &amp; note body</w:t></w:r></w:p></w:footnote>'
    '</w:footnotes>'
)


def build_edges(out_path: Path) -> None:
    doc = Document()

    # 1. Heading.
    doc.add_heading("Edges Fixture", level=1)

    # 2. Body paragraph ending with a footnote reference (id 2).
    p = doc.add_paragraph("Body sentence with a note")
    run = p.add_run()
    ref = OxmlElement("w:footnoteReference")
    ref.set(qn("w:id"), "2")
    run._r.append(ref)

    # 3. A 2x2 table with deterministic cells.
    table = doc.add_table(rows=2, cols=2)
    cells = [["Region", "Depth"], ["Benthic", "Cold"]]
    for r, row in enumerate(cells):
        for c, val in enumerate(row):
            table.rows[r].cells[c].text = val

    # 4. An inline image with deterministic alt text.
    doc.add_picture(io.BytesIO(PNG_1X1), width=Inches(1))
    docpr = doc.element.body.findall(".//" + qn("wp:docPr"))[-1]
    docpr.set("descr", "a small test image")

    # Inject the footnotes Part and relate it (no footnote author API on 1.2.0).
    part = Part(
        PackURI("/word/footnotes.xml"),
        CONTENT_TYPE.WML_FOOTNOTES,
        FOOTNOTES_XML.encode("utf-8"),
        doc.part.package,
    )
    doc.part.relate_to(part, RELATIONSHIP_TYPE.FOOTNOTES)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    edges = FIXTURES_DIR / "edges.docx"
    build_edges(edges)
    print(f"wrote {edges}")


if __name__ == "__main__":
    main()
